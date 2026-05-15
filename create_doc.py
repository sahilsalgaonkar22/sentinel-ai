"""
Sentinel AI -- Comprehensive Project Documentation (Word)
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import datetime

doc = Document()

# ---- Style Setup ----
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Headers
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Rows
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()

# ════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SENTINEL AI")
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Sovereign Observer System")
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AI-Powered Enterprise Security Intelligence Platform\nfor Automated Vulnerability Detection, Risk Analysis & Compliance Mapping")
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

for _ in range(4):
    doc.add_paragraph()

info_lines = [
    "National Forensic Sciences University",
    "M.Tech Cyber Security & Digital Forensics",
    f"Date: {datetime.date.today().strftime('%B %Y')}",
]
for line in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.size = Pt(12)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    "1. Introduction",
    "2. Problem Statement",
    "3. Objectives",
    "4. System Architecture",
    "5. Technology Stack",
    "6. Integrated Security Tools",
    "7. Key Features",
    "   7.1 Real-Time Command Center Dashboard",
    "   7.2 Multi-Tool Vulnerability Scanning",
    "   7.3 AI-Powered Risk Engine",
    "   7.4 PDF Report Generation",
    "   7.5 Compliance Mapping Engine",
    "   7.6 Scan Comparison & Drift Detection",
    "   7.7 Scheduled Recurring Scans",
    "   7.8 Alerting & Notification System",
    "   7.9 Multi-Tenant Client Portal",
    "8. System Workflow",
    "9. Database Schema & Data Models",
    "10. API Endpoints",
    "11. Frontend Architecture",
    "12. Deployment & Infrastructure",
    "13. Testing & Validation Results",
    "14. Compliance Framework Mapping",
    "15. Security Measures",
    "16. Future Scope",
    "17. Conclusion",
    "18. References",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    if not item.startswith("   "):
        for run in p.runs:
            run.bold = True

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ════════════════════════════════════════════════════════════════
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph(
    "Sentinel AI is a production-grade, AI-powered enterprise security intelligence platform designed to automate "
    "the complete vulnerability management lifecycle. The platform integrates nine real-world security scanning tools "
    "into a unified architecture that can scan networks, web applications, source code repositories, containers, and "
    "cloud infrastructure from a single interface."
)
doc.add_paragraph(
    "Unlike traditional vulnerability scanners that focus on a single attack surface, Sentinel AI provides a "
    "holistic security posture assessment by combining network scanning (Nmap), web application testing (Nikto), "
    "static code analysis (Semgrep, Bandit), container scanning (Trivy), template-based detection (Nuclei), "
    "subdomain enumeration (Subfinder), HTTP probing (httpx), and secret detection (Gitleaks) into a distributed, "
    "event-driven pipeline powered by Apache Kafka."
)
doc.add_paragraph(
    "The platform goes beyond detection by enriching findings with real-time threat intelligence from the Exploit "
    "Prediction Scoring System (EPSS) and CISA Known Exploited Vulnerabilities (KEV) catalog. An AI risk engine "
    "computes composite risk scores, while an automated compliance mapping engine maps findings to OWASP Top 10 (2021), "
    "PCI-DSS v4.0, ISO 27001:2022, and NIST CSF 2.0 frameworks. Professional PDF reports with executive summaries "
    "are generated automatically for client delivery."
)
doc.add_paragraph(
    "The entire platform is containerized using Docker Compose with 10 microservices, featuring a premium React-based "
    "frontend with real-time WebSocket updates, JWT-based multi-tenant authentication, and enterprise-grade "
    "observability through Prometheus metrics and Grafana dashboards."
)

# ════════════════════════════════════════════════════════════════
# 2. PROBLEM STATEMENT
# ════════════════════════════════════════════════════════════════
doc.add_heading('2. Problem Statement', level=1)
doc.add_paragraph(
    "Organizations today face a rapidly expanding attack surface driven by cloud adoption, IoT proliferation, and "
    "the shift to microservices architectures. The key challenges in current vulnerability management include:"
)
problems = [
    ("Fragmented Tooling", "Security teams use multiple disconnected tools (Nmap for network, Nikto for web, Trivy for containers, etc.), each with its own interface, output format, and learning curve. There is no single pane of glass to unify these results."),
    ("Manual Assessment", "Vulnerability assessment remains largely manual -- security analysts must configure tools, run scans, parse output, correlate findings, and generate reports by hand. This process is slow, error-prone, and does not scale."),
    ("Lack of Contextual Intelligence", "Raw scanner output lacks context. A CVE without EPSS probability and KEV status provides no insight into real-world exploitability, making risk prioritization guesswork."),
    ("Compliance Burden", "Enterprises must demonstrate compliance with multiple frameworks (OWASP, PCI-DSS, ISO 27001, NIST CSF). Manually mapping findings to compliance controls is time-consuming and often incomplete."),
    ("No Drift Detection", "Without historical comparison, organizations cannot track their security posture improvement over time or detect regression from new deployments."),
]
for title, desc in problems:
    p = doc.add_paragraph()
    run = p.add_run(f"{title}: ")
    run.bold = True
    p.add_run(desc)

# ════════════════════════════════════════════════════════════════
# 3. OBJECTIVES
# ════════════════════════════════════════════════════════════════
doc.add_heading('3. Objectives', level=1)
objectives = [
    "Design and develop an AI-powered, distributed security platform that automates the end-to-end vulnerability management lifecycle from target intake to report delivery.",
    "Integrate 9+ real-world security scanners (Nmap, Nikto, Semgrep, Bandit, Trivy, Nuclei, Subfinder, httpx, Gitleaks) into a unified scanning pipeline with automatic tool selection based on target type.",
    "Implement real-time threat intelligence enrichment using EPSS scores and CISA KEV data to prioritize findings based on real-world exploitability.",
    "Build an automated compliance mapping engine that maps scan findings to OWASP Top 10 (2021), PCI-DSS v4.0, ISO 27001:2022, and NIST CSF 2.0 frameworks.",
    "Develop a professional PDF report generation pipeline with executive summaries, severity breakdowns, compliance sections, and risk scores suitable for client delivery.",
    "Implement scan comparison and drift detection to track security posture changes over time.",
    "Create a production-grade, containerized deployment using Docker Compose with 10 microservices, complete with health checks, observability, and auto-recovery.",
    "Deliver a premium, real-time React frontend with WebSocket updates, multi-tenant authentication, and an intuitive dark-themed UI.",
]
for i, obj in enumerate(objectives, 1):
    doc.add_paragraph(f"{i}. {obj}")

# ════════════════════════════════════════════════════════════════
# 4. SYSTEM ARCHITECTURE
# ════════════════════════════════════════════════════════════════
doc.add_heading('4. System Architecture', level=1)
doc.add_paragraph(
    "Sentinel AI follows an event-driven microservices architecture with clear separation of concerns across five layers:"
)

doc.add_heading('4.1 Presentation Layer (Frontend)', level=2)
doc.add_paragraph(
    "A single-page application built with React 18 and Vite, served via NGINX. The frontend communicates with the "
    "API gateway through RESTful endpoints and WebSocket connections for real-time scan progress updates. The UI "
    "features a premium dark theme with Framer Motion animations, Recharts for data visualization, and Lucide React "
    "icons. Key pages include: Command Center, Threat Intelligence, Vulnerability Lab, Attack Graph, Network Map, "
    "AI Insights, Asset Cluster, Scan Management, Live Monitor, MLOps Dashboard, Analytics, Reporting, and Compliance Center."
)

doc.add_heading('4.2 API Gateway Layer', level=2)
doc.add_paragraph(
    "A FastAPI-based RESTful gateway with 15+ route groups, JWT authentication middleware, rate limiting (via slowapi), "
    "CORS configuration, and request validation. The gateway handles authentication, authorization, scan orchestration, "
    "report generation, compliance mapping, and WebSocket management. All endpoints are scoped to the authenticated "
    "user's organization (org_id) for multi-tenant isolation."
)

doc.add_heading('4.3 Message Broker Layer (Apache Kafka)', level=2)
doc.add_paragraph(
    "Apache Kafka serves as the central event bus for asynchronous scan dispatch. When a scan is initiated, the "
    "orchestrator publishes scan tasks to topic-specific queues (network-scans, web-scans, code-scans, etc.). "
    "Worker containers consume from these topics and execute the appropriate tools. A Dead Letter Queue (DLQ) "
    "captures failed tasks for automatic retry with exponential backoff, ensuring zero message loss."
)

doc.add_heading('4.4 Worker Layer (Distributed Scanners)', level=2)
doc.add_paragraph(
    "Stateless worker containers execute security scans dispatched via Kafka. Each worker runs a unified worker "
    "process that supports all 9 scanner tools. The tool_executor module manages subprocess execution with "
    "configurable timeouts, output parsing, and result normalization. After scanning, findings are enriched "
    "with EPSS scores and CISA KEV data before being stored in the database."
)

doc.add_heading('4.5 Data Layer', level=2)
doc.add_paragraph(
    "The data layer comprises five storage systems: (1) PostgreSQL as the primary relational database for scans, "
    "findings, users, and assets; (2) Redis for JWT blacklisting, caching, and rate limiting; (3) Elasticsearch "
    "for full-text search and log indexing; (4) MinIO for S3-compatible object storage of PDF reports; and "
    "(5) Prometheus + Grafana for metrics collection and operational dashboards."
)

# ════════════════════════════════════════════════════════════════
# 5. TECHNOLOGY STACK
# ════════════════════════════════════════════════════════════════
doc.add_heading('5. Technology Stack', level=1)
add_table(doc,
    ["Category", "Technology", "Purpose"],
    [
        ["Frontend Framework", "React 18 + Vite", "Single-page application with hot module reload"],
        ["UI Animation", "Framer Motion", "Smooth page transitions and micro-animations"],
        ["Charts", "Recharts", "Data visualization (severity charts, scan timelines)"],
        ["Icons", "Lucide React", "Consistent icon set across all components"],
        ["Backend Framework", "FastAPI (Python)", "Asynchronous REST API with automatic OpenAPI docs"],
        ["ORM", "SQLAlchemy 2.0 (async)", "Database abstraction with async session management"],
        ["Authentication", "JWT + bcrypt", "Stateless token auth with password hashing"],
        ["Message Broker", "Apache Kafka", "Event-driven async scan dispatch and result processing"],
        ["Primary Database", "PostgreSQL 15", "Relational storage for scans, findings, users, assets"],
        ["Caching/Rate Limiting", "Redis 7", "JWT blacklisting, API caching, rate limit counters"],
        ["Search Engine", "Elasticsearch 8", "Full-text search and log indexing"],
        ["Object Storage", "MinIO", "S3-compatible storage for PDF reports and artifacts"],
        ["Metrics", "Prometheus", "Time-series metrics collection from all services"],
        ["Dashboards", "Grafana", "Operational monitoring and alerting dashboards"],
        ["Containerization", "Docker Compose", "Multi-service orchestration with health checks"],
        ["Web Server", "NGINX", "Static asset serving and reverse proxy for frontend"],
        ["PDF Generation", "ReportLab", "Professional PDF report creation with charts"],
    ],
    [4, 5, 8]
)

# ════════════════════════════════════════════════════════════════
# 6. INTEGRATED SECURITY TOOLS
# ════════════════════════════════════════════════════════════════
doc.add_heading('6. Integrated Security Tools', level=1)
doc.add_paragraph(
    "Sentinel AI integrates nine production-grade security scanners, each automatically selected based on target type:"
)
add_table(doc,
    ["Tool", "Type", "Target", "Detection Capability"],
    [
        ["Nmap", "Network Scanner", "IP / CIDR / Domain", "Open ports, services, OS fingerprinting, NSE scripts"],
        ["Nikto", "Web Scanner", "URL / Domain", "Web server misconfigurations, outdated software, dangerous files"],
        ["Semgrep", "SAST", "Source Code", "Code vulnerabilities, injection flaws, insecure patterns"],
        ["Bandit", "Python Linter", "Python Code", "Python-specific security issues, hardcoded secrets, exec() usage"],
        ["Trivy", "Container Scanner", "Docker Image / Repo", "CVEs in OS packages, language deps, IaC misconfigurations"],
        ["Nuclei", "Template Scanner", "URL / Domain", "CVE-specific detection using 8000+ community templates"],
        ["Subfinder", "Recon", "Domain", "Subdomain enumeration across 40+ data sources"],
        ["httpx", "HTTP Prober", "URL / Domain", "HTTP status, technology detection, HSTS, security headers"],
        ["Gitleaks", "Secret Scanner", "Git Repository", "API keys, passwords, tokens, certificates in code history"],
    ],
    [3, 4, 4, 8]
)
doc.add_paragraph(
    "The input_detector module analyzes the submitted target and automatically determines which tools to execute. "
    "For example, an IP address triggers Nmap; a URL triggers Nikto, Nuclei, httpx, and Subfinder; a Git repository "
    "URL triggers Semgrep, Bandit, and Gitleaks."
)

# ════════════════════════════════════════════════════════════════
# 7. KEY FEATURES
# ════════════════════════════════════════════════════════════════
doc.add_heading('7. Key Features', level=1)

doc.add_heading('7.1 Real-Time Command Center Dashboard', level=2)
doc.add_paragraph(
    "The Command Center provides a unified overview of the organization's security posture with live metrics including: "
    "Global Risk Index (computed from finding severities), Total Vulnerabilities (Critical/High/Medium/Low breakdown), "
    "Active Scans, System Health Score (derived from service dependency status), Threat Stream (latest findings in "
    "real-time), Vulnerability Distribution chart, and Asset Intelligence. All data is sourced from the real database "
    "with zero mock or simulated values."
)

doc.add_heading('7.2 Multi-Tool Vulnerability Scanning', level=2)
doc.add_paragraph(
    "Users submit a target (IP, URL, domain, repository, or Docker image) through the Scan Management interface. "
    "The orchestrator automatically: (1) Detects the target type via the input_detector module; (2) Selects appropriate "
    "scanning tools; (3) Dispatches scan tasks to Kafka topics; (4) Workers execute tools in parallel with configurable "
    "timeouts; (5) Results are parsed, normalized, and enriched; (6) Findings are stored in PostgreSQL; (7) WebSocket "
    "notifications update the frontend in real-time."
)

doc.add_heading('7.3 AI-Powered Risk Engine', level=2)
doc.add_paragraph(
    "Each finding is enriched with contextual threat intelligence: (1) CVSS Base Score from the scanner output or NVD; "
    "(2) EPSS Score from api.first.org -- the probability of exploitation in the next 30 days; (3) CISA KEV Status -- "
    "whether the CVE is in the Known Exploited Vulnerabilities catalog. A composite risk score is computed as: "
    "risk_score = (CVSS * 0.4) + (EPSS * 100 * 0.3) + (KEV_bonus * 0.3), where KEV_bonus = 10 if the CVE is actively "
    "exploited. This ensures that actively exploited low-CVSS vulnerabilities are properly prioritized."
)

doc.add_heading('7.4 PDF Report Generation', level=2)
doc.add_paragraph(
    "The reporting module generates professional, client-ready PDF reports using ReportLab. Each report includes: "
    "Executive Summary with overall risk grade (A-F), Target and scan metadata, Severity Distribution (tabular and "
    "visual), Detailed Findings with title, description, CVSS score, evidence, and remediation steps, Compliance "
    "Mapping section showing pass/fail status across all four frameworks, and a generated SHA-256 integrity hash. "
    "Reports are stored in MinIO and accessible via the Reporting page."
)

doc.add_heading('7.5 Compliance Mapping Engine', level=2)
doc.add_paragraph(
    "The compliance mapper (backend/services/compliance/mapper.py) automatically evaluates scan findings against "
    "four major security frameworks:"
)
frameworks_detail = [
    ("OWASP Top 10 (2021)", "Maps findings to 10 categories: A01-Broken Access Control through A10-SSRF. Uses CWE ID matching (primary) and keyword-based heuristic analysis (fallback) to determine control status."),
    ("PCI-DSS v4.0", "Evaluates 11 requirements including firewall configuration, default credentials, stored cardholder data encryption, transmission security, anti-virus, secure systems, access restriction, ID assignment, physical security, network monitoring, and security testing."),
    ("ISO 27001:2022", "Checks 11 Annex A controls including information security policies, organization of infosec, HR security, asset management, access control, cryptography, physical security, operations security, communications security, system acquisition, and supplier relationships."),
    ("NIST CSF 2.0", "Validates 8 core functions: Identify (asset management), Protect (access control, data security, training), Detect (anomalies, continuous monitoring), Respond (planning, communications), and Recover (planning, improvements)."),
]
for fw, desc in frameworks_detail:
    p = doc.add_paragraph()
    run = p.add_run(f"{fw}: ")
    run.bold = True
    p.add_run(desc)

doc.add_heading('7.6 Scan Comparison & Drift Detection', level=2)
doc.add_paragraph(
    "The scan_comparator module enables before-vs-after analysis of two scans against the same target. The comparison "
    "output includes: number of new vulnerabilities introduced, number of resolved vulnerabilities, security score "
    "delta, severity distribution changes, and a natural-language summary (e.g., 'Security posture improved. Score "
    "increased by 5 points. 1 vulnerability resolved. 1 new vulnerability introduced.'). This enables organizations "
    "to track security posture improvement over time and detect regression from new deployments."
)

doc.add_heading('7.7 Scheduled Recurring Scans', level=2)
doc.add_paragraph(
    "The scheduler module (scheduler.py) supports cron-based recurring scans. Users can schedule scans via the "
    "API endpoint POST /scans/{scan_id}/schedule with a cron expression (e.g., '0 0 * * 1' for weekly Monday "
    "midnight). The scheduler automatically triggers new scans at the specified intervals and can integrate with "
    "the alerting system to notify on critical findings."
)

doc.add_heading('7.8 Alerting & Notification System', level=2)
doc.add_paragraph(
    "The alerting module (alerting/__init__.py) provides multi-channel notifications: (1) SMTP Email -- configurable "
    "SMTP server with TLS support for sending scan reports and critical finding alerts; (2) Slack Webhooks -- "
    "integration with Slack incoming webhooks for real-time team notifications; (3) Redis-backed rate limiting "
    "to prevent alert storms during large scans. Alert triggers include: critical severity findings, KEV-flagged "
    "CVEs, scan completion, and system health degradation."
)

doc.add_heading('7.9 Multi-Tenant Client Portal', level=2)
doc.add_paragraph(
    "All data access is scoped to the authenticated user's organization (org_id). This ensures: (1) Client A cannot "
    "see Client B's scans, findings, or reports; (2) Each organization has independent scan histories, assets, and "
    "compliance status; (3) JWT tokens encode the org_id, and all database queries filter by it; (4) Role-based "
    "access control (RBAC) can be extended to support admin, analyst, and viewer roles."
)

# ════════════════════════════════════════════════════════════════
# 8. SYSTEM WORKFLOW
# ════════════════════════════════════════════════════════════════
doc.add_heading('8. System Workflow', level=1)
doc.add_paragraph("The end-to-end scan workflow follows these steps:")
steps = [
    "User Authentication: The user logs in via the Security Gate (login page) with email and password. The gateway validates credentials against bcrypt-hashed passwords in PostgreSQL and issues a JWT token.",
    "Target Submission: The user navigates to Scan Management and enters a target (e.g., 'kartik-rathi.site'). The frontend sends a POST /scans/ request with the target and optional scan configuration.",
    "Input Detection: The input_detector module analyzes the target string and classifies it as IP, URL, domain, repository, or Docker image.",
    "Tool Selection: Based on the target type, the orchestrator selects the appropriate scanning tools from the 9 available scanners.",
    "Kafka Dispatch: The orchestrator publishes scan tasks to the corresponding Kafka topics (network-scans, web-scans, code-scans). Each message contains the scan ID, target, and tool configuration.",
    "Worker Execution: Worker containers consume messages from Kafka and execute the selected tools as subprocesses. Each tool runs with configurable timeouts and output parsing.",
    "Result Processing: Raw scanner output is parsed into normalized Finding objects with standardized fields (title, severity, CVE, CWE, CVSS, evidence, remediation).",
    "Threat Intelligence Enrichment: Each finding with a CVE ID is enriched with EPSS score (from api.first.org) and CISA KEV status (from cisa.gov). A composite risk score is computed.",
    "Database Storage: Normalized findings are stored in PostgreSQL with full audit trail (scan_id, org_id, timestamps).",
    "Real-Time Update: WebSocket messages notify the frontend of scan progress and completion. The dashboard updates in real-time.",
    "Report Generation: The user can generate a PDF report from the Reporting page. The pdf_generator module creates a professional report with executive summary, findings, and compliance mapping.",
    "Compliance Mapping: The compliance mapper evaluates findings against OWASP, PCI-DSS, ISO 27001, and NIST CSF controls, producing a compliance score and detailed control-level status.",
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"Step {i} -- ")
    run.bold = True
    p.add_run(step)

# ════════════════════════════════════════════════════════════════
# 9. DATABASE SCHEMA
# ════════════════════════════════════════════════════════════════
doc.add_heading('9. Database Schema & Data Models', level=1)
doc.add_paragraph(
    "The PostgreSQL database contains the following core tables, defined in backend/services/scan_control/models.py:"
)
add_table(doc,
    ["Model", "Key Fields", "Purpose"],
    [
        ["User", "id, email, hashed_password, org_id, role, is_active, created_at", "User accounts with org-scoped multi-tenancy"],
        ["Scan", "id, name, target_raw, status, security_score, tools_used, org_id, created_at, completed_at", "Scan metadata and lifecycle tracking"],
        ["Finding", "id, scan_id, title, severity, description, cve_id, cwe_id, cvss_score, epss_score, tool_name, evidence, remediation, org_id", "Individual vulnerability findings enriched with threat intel"],
        ["Asset", "id, name, type, risk_score, org_id, created_at", "Discovered assets (IPs, domains, applications)"],
        ["Vulnerability", "id, title, severity, risk_score, org_id", "Aggregated vulnerability catalog (legacy)"],
        ["AttackPath", "id, name, severity, risk_score, chain_steps, org_id", "AI-computed attack path chains"],
    ],
    [3, 8, 6]
)

# ════════════════════════════════════════════════════════════════
# 10. API ENDPOINTS
# ════════════════════════════════════════════════════════════════
doc.add_heading('10. API Endpoints', level=1)
doc.add_paragraph(
    "The FastAPI gateway exposes the following route groups, all requiring JWT authentication (except /auth/login and /auth/register):"
)
add_table(doc,
    ["Route Group", "Key Endpoints", "Description"],
    [
        ["/auth", "POST /login, POST /register, GET /me", "Authentication and user management"],
        ["/scans", "GET /, POST /, GET /{id}, GET /{id}/findings, POST /compare, POST /{id}/schedule", "Scan lifecycle management, comparison, scheduling"],
        ["/dashboard", "GET /stats, GET /analytics", "Aggregated dashboard statistics and analytics"],
        ["/vulnerabilities", "GET /, GET /stats/summary, GET /attack-paths/, GET /{id}", "Vulnerability listing and detail"],
        ["/reporting", "GET /, POST /generate, GET /download/{name}", "PDF report generation and retrieval"],
        ["/compliance", "GET /{scan_id}, GET /frameworks/list", "Compliance mapping and framework listing"],
        ["/assets", "GET /, POST /, GET /{id}", "Asset inventory management"],
        ["/ai", "GET /metrics, GET /attack-graph", "AI engine metrics and attack graph data"],
        ["/settings", "GET /, PUT /", "Platform configuration"],
        ["/alerts", "GET /, POST /webhook", "Alert configuration and webhook management"],
        ["/live-scan", "GET /tools, WS /ws/{scan_id}", "Live scan monitoring and WebSocket streams"],
        ["/health", "GET /", "System health check with dependency status"],
    ],
    [3, 7, 7]
)

# ════════════════════════════════════════════════════════════════
# 11. FRONTEND ARCHITECTURE
# ════════════════════════════════════════════════════════════════
doc.add_heading('11. Frontend Architecture', level=1)
doc.add_paragraph(
    "The frontend is a React 18 single-page application built with Vite, featuring 13 distinct pages accessible "
    "via a persistent sidebar navigation:"
)
add_table(doc,
    ["Page", "Route", "Description"],
    [
        ["Command Center", "/", "Central dashboard with risk index, threat stream, vulnerability distribution"],
        ["Threat Intelligence", "/threat-intel", "Searchable findings list with severity filtering and detail panel"],
        ["Vulnerability Lab", "/vulnerabilities", "Paginated vulnerability database with CVSS and CVE data"],
        ["Attack Graph", "/attack-graph", "Canvas-based force-directed graph of vulnerabilities, assets, and actors"],
        ["Network Map", "/network", "Network topology visualization"],
        ["AI Insights", "/ai-insights", "AI engine analysis and recommendations"],
        ["Asset Cluster", "/assets", "Asset inventory with risk scoring"],
        ["Scan Management", "/scans", "Scan initiation, history, and status tracking"],
        ["Live Monitor", "/live-monitor", "Real-time scan progress via WebSocket"],
        ["MLOps Dashboard", "/ai-monitor", "AI model metrics and retraining status"],
        ["Analytics", "/analytics", "Historical scan data, severity trends, and charts"],
        ["Reporting", "/reporting", "PDF report generation, download, and management"],
        ["Compliance Center", "/compliance", "Compliance mapping across 4 frameworks with drill-down"],
    ],
    [4, 4, 9]
)
doc.add_paragraph(
    "The frontend uses a centralized API client (api/client.js) with interceptors for JWT token management, "
    "a Zustand store (authStore.js) for authentication state, and Framer Motion for page transitions."
)

# ════════════════════════════════════════════════════════════════
# 12. DEPLOYMENT & INFRASTRUCTURE
# ════════════════════════════════════════════════════════════════
doc.add_heading('12. Deployment & Infrastructure', level=1)
doc.add_paragraph(
    "Sentinel AI is deployed using Docker Compose with 10 containers, all with health checks and auto-restart policies:"
)
add_table(doc,
    ["Container", "Image", "Port", "Health Check"],
    [
        ["sentinel-frontend", "NGINX + React build", "3000", "HTTP /health endpoint"],
        ["sentinel-gateway", "Python FastAPI", "8000", "HTTP /health with dependency checks"],
        ["sentinel-worker", "Python + 9 scanners", "--", "Custom health check script"],
        ["sentinel-postgres", "PostgreSQL 15", "5432", "pg_isready"],
        ["sentinel-redis", "Redis 7", "6379", "redis-cli ping"],
        ["sentinel-kafka", "Confluent Kafka", "9092, 29092", "kafka-broker-api-versions"],
        ["sentinel-elasticsearch", "Elasticsearch 8", "9200", "HTTP cluster health"],
        ["sentinel-minio", "MinIO", "9000, 9001", "HTTP /minio/health/live"],
        ["sentinel-prometheus", "Prometheus", "9090", "--"],
        ["sentinel-grafana", "Grafana", "3001", "--"],
    ],
    [4, 4, 3, 5]
)

# ════════════════════════════════════════════════════════════════
# 13. TESTING & VALIDATION
# ════════════════════════════════════════════════════════════════
doc.add_heading('13. Testing & Validation Results', level=1)
doc.add_paragraph(
    "A comprehensive pre-presentation audit was conducted to validate all system functionality:"
)

doc.add_heading('13.1 API Endpoint Testing', level=2)
doc.add_paragraph(
    "An automated audit script tested all 43 API endpoints across 12 categories. Results: 43/43 PASSED (100% pass rate). "
    "This includes authentication, scan management, scan comparison, compliance mapping, PDF generation, dashboard "
    "analytics, vulnerability listing, asset management, settings, and live monitoring."
)

doc.add_heading('13.2 Infrastructure Health', level=2)
doc.add_paragraph(
    "All 10 Docker containers were verified as running and healthy. PostgreSQL, Redis, Kafka, Elasticsearch, and MinIO "
    "all passed their respective health checks. The system health score reported by the gateway was 100%."
)

doc.add_heading('13.3 Real Scan Validation', level=2)
doc.add_paragraph(
    "A live scan was executed against 'kartik-rathi.site' using the full tool pipeline. The scan completed successfully, "
    "detecting 5 real vulnerabilities (3 Medium, 2 Low) including Missing Content-Security-Policy, Missing X-Frame-Options, "
    "and Technology Stack Detection (HSTS, Vercel). All findings were enriched with EPSS scores and stored in the database."
)

doc.add_heading('13.4 Compliance Validation', level=2)
doc.add_paragraph(
    "The compliance mapping engine was validated against the scan results, producing: OWASP Top 10: 8/10 controls passed "
    "(2 violations), PCI-DSS v4.0: 10/11 controls passed (1 violation), ISO 27001:2022: 11/11 controls passed "
    "(fully compliant), NIST CSF 2.0: 7/8 controls passed (1 gap). Overall compliance score: 90%."
)

doc.add_heading('13.5 Report Generation', level=2)
doc.add_paragraph(
    "PDF report generation was tested successfully, producing a 6.1 KB professional report with executive summary, "
    "severity breakdown, individual findings, and compliance mapping section. The report was stored in MinIO with "
    "SHA-256 integrity hash."
)

# ════════════════════════════════════════════════════════════════
# 14. COMPLIANCE FRAMEWORK MAPPING
# ════════════════════════════════════════════════════════════════
doc.add_heading('14. Compliance Framework Mapping', level=1)
doc.add_paragraph(
    "The compliance mapping engine (mapper.py) uses a dual-detection strategy:"
)
doc.add_paragraph(
    "Primary Detection (CWE-based): Each framework control is associated with a list of CWE IDs. When a finding "
    "has a matching CWE ID, the control is marked as FAIL."
)
doc.add_paragraph(
    "Secondary Detection (Keyword-based): For findings without CWE data, keyword matching is applied against the "
    "finding title and description. Keywords are mapped to specific controls (e.g., 'SQL injection' maps to OWASP A03, "
    "'missing encryption' maps to PCI-DSS Requirement 4)."
)
add_table(doc,
    ["Framework", "Controls Evaluated", "Sample Pass Result", "Sample Fail Result"],
    [
        ["OWASP Top 10", "10 (A01-A10)", "A04: Insecure Design - PASS", "A05: Security Misconfiguration - FAIL (Missing CSP header)"],
        ["PCI-DSS v4.0", "11 (Req 1-11)", "Req 3: Protect Stored Data - PASS", "Req 6: Secure Systems - FAIL (Missing X-Frame-Options)"],
        ["ISO 27001", "11 (Annex A)", "A.10: Cryptography - PASS", "All controls passed"],
        ["NIST CSF", "8 (ID/PR/DE/RS/RC)", "PR.AC: Access Control - PASS", "DE.CM: Continuous Monitoring - FAIL"],
    ],
    [3, 4, 5, 6]
)

# ════════════════════════════════════════════════════════════════
# 15. SECURITY MEASURES
# ════════════════════════════════════════════════════════════════
doc.add_heading('15. Security Measures', level=1)
measures = [
    ("Authentication", "JWT tokens with configurable expiry, bcrypt password hashing with salt rounds, token blacklisting via Redis for logout."),
    ("Authorization", "Organization-scoped access control (org_id). All database queries filter by the authenticated user's organization."),
    ("Rate Limiting", "API rate limiting via slowapi (60 requests/minute per endpoint) with Redis-backed counters."),
    ("Input Validation", "Pydantic schema validation on all API inputs. Scope validation prevents scanning unauthorized targets."),
    ("Transport Security", "HTTPS-ready NGINX configuration. CORS whitelist for frontend origin only."),
    ("Secret Management", "Passwords and API keys stored as environment variables via .env files, never hardcoded."),
    ("Container Security", "Non-root container execution, minimal base images, health checks, auto-restart policies."),
    ("Audit Logging", "Structured JSON logging with timestamps, log levels, module names, and request correlation IDs."),
]
for title, desc in measures:
    p = doc.add_paragraph()
    run = p.add_run(f"{title}: ")
    run.bold = True
    p.add_run(desc)

# ════════════════════════════════════════════════════════════════
# 16. FUTURE SCOPE
# ════════════════════════════════════════════════════════════════
doc.add_heading('16. Future Scope', level=1)
future = [
    ("LLM-Powered Remediation", "Integration with GPT-4/Gemini to provide finding-specific remediation guidance with code fix suggestions."),
    ("Kubernetes Deployment", "Migration from Docker Compose to Kubernetes with auto-scaling worker pods based on scan queue depth."),
    ("SIEM Integration", "Connectors for Splunk, QRadar, and Azure Sentinel for enterprise SOC workflow automation."),
    ("Mobile Companion App", "React Native mobile application with push notifications for critical findings."),
    ("MITRE ATT&CK Mapping", "Map findings to MITRE ATT&CK tactics and techniques for threat-informed defense."),
    ("Threat Intel Feeds", "Integration with AlienVault OTX, VirusTotal, and Shodan for enriched intelligence."),
    ("Custom Scanner Plugins", "Plugin architecture allowing organizations to integrate proprietary or custom scanning tools."),
    ("CI/CD Pipeline Integration", "GitHub Actions / GitLab CI plugins for shift-left security in development workflows."),
]
for title, desc in future:
    p = doc.add_paragraph()
    run = p.add_run(f"{title}: ")
    run.bold = True
    p.add_run(desc)

# ════════════════════════════════════════════════════════════════
# 17. CONCLUSION
# ════════════════════════════════════════════════════════════════
doc.add_heading('17. Conclusion', level=1)
doc.add_paragraph(
    "Sentinel AI demonstrates the feasibility and value of building an AI-powered, end-to-end security intelligence "
    "platform that automates the complete vulnerability management lifecycle. The platform successfully integrates "
    "nine production-grade security scanners into a distributed, event-driven architecture capable of scanning "
    "networks, web applications, source code repositories, containers, and cloud infrastructure from a single interface."
)
doc.add_paragraph(
    "Key achievements include: (1) Real vulnerability detection validated on live targets with enriched threat "
    "intelligence (EPSS + CISA KEV); (2) Automated compliance mapping achieving 90% compliance score across four "
    "major frameworks (OWASP, PCI-DSS, ISO 27001, NIST CSF); (3) Professional PDF report generation suitable for "
    "client delivery; (4) Enterprise-grade features including multi-tenancy, scheduled scans, drift detection, and "
    "multi-channel alerting; (5) A production-ready containerized deployment with 10 microservices, 100% health "
    "check pass rate, and zero mock data."
)
doc.add_paragraph(
    "The platform validates the thesis that AI-augmented automation can significantly reduce the time, cost, and "
    "error rate of enterprise vulnerability management while providing actionable intelligence for security teams "
    "and compliance-ready documentation for management."
)

# ════════════════════════════════════════════════════════════════
# 18. REFERENCES
# ════════════════════════════════════════════════════════════════
doc.add_heading('18. References', level=1)
refs = [
    "OWASP Foundation. (2021). OWASP Top Ten Web Application Security Risks. https://owasp.org/Top10/",
    "PCI Security Standards Council. (2022). PCI DSS v4.0. https://www.pcisecuritystandards.org/",
    "ISO/IEC. (2022). ISO 27001:2022 Information Security Management Systems. https://www.iso.org/standard/27001",
    "NIST. (2024). Cybersecurity Framework 2.0. https://www.nist.gov/cyberframework",
    "FIRST.org. (2023). Exploit Prediction Scoring System (EPSS). https://api.first.org/data/v1/epss",
    "CISA. (2024). Known Exploited Vulnerabilities Catalog. https://www.cisa.gov/known-exploited-vulnerabilities-catalog",
    "Nmap Project. Network Mapper. https://nmap.org/",
    "Sullo, C. Nikto Web Server Scanner. https://github.com/sullo/nikto",
    "Semgrep. Static Analysis at Scale. https://semgrep.dev/",
    "Bandit. Python Security Linter. https://bandit.readthedocs.io/",
    "Aqua Security. Trivy Vulnerability Scanner. https://trivy.dev/",
    "ProjectDiscovery. Nuclei Template-Based Scanner. https://nuclei.projectdiscovery.io/",
    "FastAPI. Modern Python Web Framework. https://fastapi.tiangolo.com/",
    "Apache Kafka. Distributed Event Streaming Platform. https://kafka.apache.org/",
    "React. A JavaScript Library for User Interfaces. https://react.dev/",
]
for i, ref in enumerate(refs, 1):
    doc.add_paragraph(f"[{i}] {ref}")

# ════════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════════
output_path = r"C:\Users\SAHIL\Downloads\ransom2\Sentinel_AI_Project_Report.docx"
doc.save(output_path)
print("Saved to:", output_path)
